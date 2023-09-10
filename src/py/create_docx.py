import datetime
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
from docx.shared import RGBColor
from pathlib import Path

SPACING = 32
DOSYA_ISMI = f'{Path.home()}/Dev/zimmet/output/zimmet.docx'
LOGO_IMG = f'{Path.home()}/Dev/zimmet/assets/mega.png'

# daha kolay olsun diye object olustur
class Asset:
  def __init__(self, loaded_json):
    raw_date = loaded_json.get('last_checkout').get('datetime')
    datetime_object = datetime.datetime.strptime(raw_date, '%Y-%m-%d %H:%M:%S')
    self.last_checkout = datetime_object.strftime("%d.%m.%Y %H:%M:%S")
    self.manufacturer = loaded_json.get('manufacturer').get('name')
    self.model = loaded_json.get('model').get('name')
    self.category = loaded_json.get('category').get('name')
    self.serial = loaded_json.get('serial')

# tarih: str, marka: str, model: str, kategori: str, seri: str
def create_single_asset_docx(asset: Asset):
    # asil dokuman
    document = Document()

    # paragraf bosluklari
    # paragraph_format = document.styles['Normal'].paragraph_format
    # paragraph_format.space_before = Pt(81)

    # logo tarih table
    table = document.add_table(rows=1, cols=2)

    # resim ekle
    pic_cel = table.rows[0].cells[0].paragraphs[0]
    run = pic_cel.add_run()
    run.add_picture(LOGO_IMG, width=Inches(1.25))

    # bugunun tarihi
    row = table.rows[0].cells
    row[1].text = asset.last_checkout

    # table ortala
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # baslik
    heading = document.add_heading('Zimmet Teslim Tutanağı', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(SPACING)

    # text
    zimmet_paragraf = document.add_paragraph(f'"{asset.manufacturer}" marka "{asset.model}" model, "{asset.category}" kategorisindeki "{asset.serial}" seri numaralı cihaz çalışır bir şekilde teslim edilmiştir.')
    zimmet_paragraf.paragraph_format.space_before = Pt(SPACING)
    zimmet_paragraf.paragraph_format.space_after = Pt(SPACING)

    # isimlerin yazacagi
    isimler = document.add_table(rows=1, cols=2)
    isimler_row = isimler.rows[0].cells
    isimler_row[0].text = "Teslim Eden"
    isimler_row[1].text = "Teslim Alan"
    isimler_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    isimler_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    isimler_row[0].paragraphs[0].paragraph_format.space_before = Pt(SPACING*2)
    isimler_row[1].paragraphs[0].paragraph_format.space_before = Pt(SPACING*2)

    # dosyayi olustur
    document.save(DOSYA_ISMI)

def create_user_assets_docx(assets: [Asset]):
    # asil dokuman
    document = Document()

    # paragraf bosluklari
    # paragraph_format = document.styles['Normal'].paragraph_format
    # paragraph_format.space_before = Pt(81)

    # logo tarih table
    table = document.add_table(rows=1, cols=2)

    # resim ekle
    pic_cel = table.rows[0].cells[0].paragraphs[0]
    run = pic_cel.add_run()
    run.add_picture(LOGO_IMG, width=Inches(1.25))

    # bugunun tarihi
    row = table.rows[0].cells
    row[1].text = assets[0].last_checkout

    # table ortala
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # baslik
    heading = document.add_heading('Zimmet Teslim Tutanağı', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(SPACING)

    # text
    # zimmet_paragraf = document.add_paragraph(f'"{asset.manufacturer}" marka "{asset.model}" model, "{asset.category}" kategorisindeki "{asset.serial}" seri numaralı cihaz çalışır bir şekilde teslim edilmiştir.')
    zimmet_paragraf = document.add_paragraph("Aşağıda yer alan ürünler çalışır şekilde teslim edilmiştir.")
    zimmet_paragraf.paragraph_format.space_before = Pt(SPACING)
    zimmet_paragraf.paragraph_format.space_after = Pt(SPACING)

    # demirbas listesi ----------------------------------------------------------------
    # basliklar
    asset_table = document.add_table(rows=1, cols=5)
    asset_header_row = asset_table.rows[0]
    asset_header_row.cells[0].text = "Marka"
    asset_header_row.cells[1].text = "Model"
    asset_header_row.cells[2].text = "Kategori"
    asset_header_row.cells[3].text = "Seri No"
    asset_header_row.cells[4].text = "Tarih"

    # urunler
    for asset in assets:
        asset_row = asset_table.add_row().cells
        asset_row[0].text = asset.manufacturer
        asset_row[1].text = asset.model
        asset_row[2].text = asset.category
        asset_row[3].text = asset.serial
        asset_row[4].text = asset.last_checkout

    asset_table.style = 'Light Shading Accent 1'

    # demirbas listesi ----------------------------------------------------------------

    # isimlerin yazacagi
    isimler = document.add_table(rows=1, cols=2)
    isimler_row = isimler.rows[0].cells
    isimler_row[0].text = "Teslim Eden"
    isimler_row[1].text = "Teslim Alan"
    isimler_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    isimler_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    isimler_row[0].paragraphs[0].paragraph_format.space_before = Pt(SPACING*2)
    isimler_row[1].paragraphs[0].paragraph_format.space_before = Pt(SPACING*2)

    # dosyayi olustur
    document.save(DOSYA_ISMI)